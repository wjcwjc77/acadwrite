"""
输出生成模块，负责生成最终文档
"""

import os
import logging
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)


class OutputGenerator(ABC):
    """
    输出生成器基类，定义通用接口
    """
    
    @abstractmethod
    def generate(self, styled_content: Dict[str, Any], template_structure: Dict[str, Any], output_file: str) -> None:
        """
        生成输出文档
        
        Args:
            styled_content: 应用样式后的内容结构
            template_structure: 模板结构
            output_file: 输出文件路径
        """
        pass


class DocxOutputGenerator(OutputGenerator):
    """
    .docx格式输出生成器
    """
    
    def generate(self, styled_content: Dict[str, Any], template_structure: Dict[str, Any], output_file: str) -> None:
        """
        生成.docx输出文档
        
        Args:
            styled_content: 应用样式后的内容结构
            template_structure: 模板结构
            output_file: 输出文件路径
        """
        logger.info(f"开始生成.docx输出文档: {output_file}")
        
        try:
            # 使用模板创建新文档
            template_path = template_structure.get('path', '')
            if os.path.exists(template_path):
                doc = docx.Document(template_path)
                # 清除模板中的内容，但保留样式
                for i in range(len(doc.paragraphs)-1, -1, -1):
                    p = doc.paragraphs[i]
                    p._element.getparent().remove(p._element)
            else:
                logger.warning(f"模板文件不存在: {template_path}，使用空白文档")
                doc = docx.Document()
            
            # 处理每个内容元素
            elements = styled_content.get('elements', [])
            for element in elements:
                self._add_element_to_doc(doc, element)
            
            # 保存文档
            doc.save(output_file)
            logger.info(f".docx输出文档生成完成: {output_file}")
            
        except Exception as e:
            logger.error(f"生成.docx输出文档时出错: {str(e)}", exc_info=True)
            raise
    
    def _add_element_to_doc(self, doc: docx.Document, element: Dict[str, Any]) -> None:
        """
        将元素添加到文档中
        
        Args:
            doc: docx文档对象
            element: 内容元素
        """
        element_type = element.get('type', '')
        
        if element_type == 'heading':
            # 添加标题
            level = element.get('level', 1)
            text = element.get('text', '')
            style = element.get('style', f'Heading {level}')
            
            paragraph = doc.add_paragraph(text, style=style)
            
        elif element_type == 'paragraph':
            # 添加段落
            text = element.get('text', '')
            style = element.get('style', 'Normal')
            
            paragraph = doc.add_paragraph(text, style=style)
            
        elif element_type == 'list_item':
            # 添加列表项
            text = element.get('text', '')
            style = element.get('style', 'List Paragraph')
            list_type = element.get('list_type', 'unordered')
            
            paragraph = doc.add_paragraph(text, style=style)
            paragraph.style = style
            
            # 设置列表格式
            if list_type == 'unordered':
                paragraph.style.paragraph_format.left_indent = Inches(0.25)
                paragraph.style.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                paragraph.style.paragraph_format.left_indent = Inches(0.25)
                paragraph.style.paragraph_format.first_line_indent = Inches(-0.25)
            
        elif element_type == 'code_block':
            # 添加代码块
            text = element.get('text', '')
            style = element.get('style', 'Code')
            
            paragraph = doc.add_paragraph(text, style=style)
            
        elif element_type == 'block_quote':
            # 添加引用块
            text = element.get('text', '')
            style = element.get('style', 'Quote')
            
            paragraph = doc.add_paragraph(text, style=style)
            
        elif element_type == 'table':
            # 添加表格
            rows = element.get('rows', [])
            # 简化处理，实际应用中需要更复杂的表格处理逻辑
            table = doc.add_table(rows=1, cols=1)
            table.style = element.get('style', 'Table Normal')
            
        elif element_type == 'image':
            # 添加图片
            src = element.get('src', '')
            alt = element.get('alt', '')
            
            # 简化处理，实际应用中需要处理图片路径和插入逻辑
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 如果有图片路径且文件存在，则添加图片
            if src and os.path.exists(src):
                run = paragraph.add_run()
                run.add_picture(src)
                
                # 添加图片说明
                if alt:
                    caption = doc.add_paragraph(alt, style=element.get('caption_style', 'Caption'))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


class TexOutputGenerator(OutputGenerator):
    """
    .tex格式输出生成器
    """
    
    def generate(self, styled_content: Dict[str, Any], template_structure: Dict[str, Any], output_file: str) -> None:
        """
        生成.tex输出文档
        
        Args:
            styled_content: 应用样式后的内容结构
            template_structure: 模板结构
            output_file: 输出文件路径
        """
        logger.info(f"开始生成.tex输出文档: {output_file}")
        
        try:
            # 获取模板内容
            template_path = template_structure.get('path', '')
            template_content = ''
            if os.path.exists(template_path):
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_content = f.read()
            
            # 提取文档前导部分和后续部分
            preamble, document_env = self._extract_tex_parts(template_content)
            
            # 生成内容部分
            content = self._generate_content(styled_content)
            
            # 组合最终文档
            if document_env:
                # 使用模板的文档环境
                begin_pos = document_env.find('\\begin{document}') + len('\\begin{document}')
                end_pos = document_env.rfind('\\end{document}')
                
                if begin_pos >= 0 and end_pos >= 0:
                    output_content = document_env[:begin_pos] + '\n\n' + content + '\n\n' + document_env[end_pos:]
                else:
                    output_content = preamble + '\n\\begin{document}\n\n' + content + '\n\n\\end{document}\n'
            else:
                # 创建新的文档环境
                output_content = preamble + '\n\\begin{document}\n\n' + content + '\n\n\\end{document}\n'
            
            # 保存文档
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(output_content)
            
            logger.info(f".tex输出文档生成完成: {output_file}")
            
        except Exception as e:
            logger.error(f"生成.tex输出文档时出错: {str(e)}", exc_info=True)
            raise
    
    def _extract_tex_parts(self, template_content: str) -> tuple:
        """
        从模板内容中提取前导部分和文档环境
        
        Args:
            template_content: 模板内容
            
        Returns:
            前导部分和文档环境
        """
        # 查找文档环境
        begin_pos = template_content.find('\\begin{document}')
        end_pos = template_content.rfind('\\end{document}') + len('\\end{document}')
        
        if begin_pos >= 0 and end_pos >= 0:
            preamble = template_content[:begin_pos]
            document_env = template_content[begin_pos:end_pos]
            return preamble, document_env
        else:
            # 没有找到文档环境，返回整个内容作为前导部分
            return template_content, ''
    
    def _generate_content(self, styled_content: Dict[str, Any]) -> str:
        """
        生成LaTeX内容
        
        Args:
            styled_content: 应用样式后的内容结构
            
        Returns:
            LaTeX内容文本
        """
        content_lines = []
        
        # 处理每个内容元素
        elements = styled_content.get('elements', [])
        for element in elements:
            element_type = element.get('type', '')
            
            if element_type == 'heading':
                # 添加标题
                level = element.get('level', 1)
                text = element.get('text', '')
                command = element.get('command', self._get_default_tex_heading_command(level))
                
                content_lines.append(f"{command}{{{text}}}")
                
            elif element_type == 'paragraph':
                # 添加段落
                text = element.get('text', '')
                content_lines.append(text)
                content_lines.append('')  # 空行分隔段落
                
            elif element_type == 'environment':
                # 添加环境
                env_type = element.get('env_type', '')
                content = element.get('content', '')
                begin_def = element.get('begin_def', f'\\begin{{{env_type}}}')
                end_def = element.get('end_def', f'\\end{{{env_type}}}')
                
                content_lines.append(begin_def)
                
                if env_type in ['itemize', 'enumerate'] and 'items' in element:
                    # 处理列表项
                    for item in element.get('items', []):
                        content_lines.append(f"\\item {item}")
                else:
                    # 处理其他环境内容
                    content_lines.append(content)
                
                content_lines.append(end_def)
                
            elif element_type == 'command':
                # 添加命令
                name = element.get('name', '')
                options = element.get('options', [])
                arguments = element.get('arguments', [])
                
                command = f"\\{name}"
                
                # 添加选项
                if options:
                    options_str = ','.join(options)
                    command += f"[{options_str}]"
                
                # 添加参数
                for arg in arguments:
                    command += f"{{{arg}}}"
                
                content_lines.append(command)
                
                # 处理图片说明
                if name == 'includegraphics' and 'caption' in element:
                    caption = element.get('caption', '')
                    if caption:
                        content_lines.append(f"\\caption{{{caption}}}")
        
        return '\n'.join(content_lines)
    
    def _get_default_tex_heading_command(self, level: int) -> str:
        """
        获取默认的LaTeX标题命令
        
        Args:
            level: 标题级别
            
        Returns:
            默认的LaTeX命令
        """
        level_to_command = {
            1: '\\section',
            2: '\\subsection',
            3: '\\subsubsection',
            4: '\\paragraph',
            5: '\\subparagraph',
            6: '\\subparagraph'
        }
        
        return level_to_command.get(level, '\\section')


class OutputGeneratorFactory:
    """
    输出生成器工厂，根据格式类型创建相应的生成器
    """
    
    @staticmethod
    def create_generator(format_type: str) -> OutputGenerator:
        """
        创建输出生成器
        
        Args:
            format_type: 输出格式类型
            
        Returns:
            对应格式的输出生成器
        """
        if format_type.lower() == 'docx':
            return DocxOutputGenerator()
        elif format_type.lower() == 'tex':
            return TexOutputGenerator()
        else:
            raise ValueError(f"不支持的输出格式: {format_type}")
