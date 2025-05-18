"""
模板解析模块，负责解析不同格式的模板文件
"""

import os
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional
import docx
from docx.document import Document as DocxDocument
from docx.styles.style import _ParagraphStyle
import re
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)


class TemplateParser(ABC):
    """
    模板解析器基类，定义通用接口
    """
    
    @abstractmethod
    def parse(self, template_file: str) -> Dict[str, Any]:
        """
        解析模板文件，返回结构化表示
        
        Args:
            template_file: 模板文件路径
            
        Returns:
            结构化的模板表示
        """
        pass


class DocxTemplateParser(TemplateParser):
    """
    .docx格式模板解析器
    """
    
    def parse(self, template_file: str) -> Dict[str, Any]:
        """
        解析.docx模板文件
        
        Args:
            template_file: .docx模板文件路径
            
        Returns:
            结构化的模板表示
        """
        logger.info(f"解析.docx模板: {template_file}")
        
        try:
            doc = docx.Document(template_file)
            
            # 提取样式信息
            styles = self._extract_styles(doc)
            
            # 提取文档结构
            structure = self._extract_structure(doc)
            
            # 提取页面设置
            page_settings = self._extract_page_settings(doc)
            
            template_info = {
                'type': 'docx',
                'path': template_file,
                'styles': styles,
                'structure': structure,
                'page_settings': page_settings
            }
            
            logger.info(f"解析完成，提取了 {len(styles)} 个样式和 {len(structure)} 个结构元素")
            return template_info
            
        except Exception as e:
            logger.error(f"解析.docx模板时出错: {str(e)}", exc_info=True)
            raise
    
    def _extract_styles(self, doc: DocxDocument) -> Dict[str, Dict[str, Any]]:
        """
        提取文档中的样式定义
        
        Args:
            doc: docx文档对象
            
        Returns:
            样式定义字典
        """
        styles = {}
        
        # 修复：使用doc.styles的名称列表而不是直接遍历style_id
        for style_name in doc.styles.element.xpath('//w:style/@w:styleId'):
            try:
                style = doc.styles[style_name]
                if isinstance(style, _ParagraphStyle):
                    style_info = {
                        'name': style.name,
                        'font': {},
                        'paragraph': {}
                    }
                    
                    # 提取字体信息
                    if hasattr(style, 'font') and style.font:
                        if hasattr(style.font, 'name') and style.font.name:
                            style_info['font']['name'] = style.font.name
                        if hasattr(style.font, 'size') and style.font.size:
                            style_info['font']['size'] = style.font.size
                        if hasattr(style.font, 'bold') and style.font.bold:
                            style_info['font']['bold'] = style.font.bold
                        if hasattr(style.font, 'italic') and style.font.italic:
                            style_info['font']['italic'] = style.font.italic
                    
                    # 提取段落信息
                    if hasattr(style, 'paragraph_format') and style.paragraph_format:
                        pf = style.paragraph_format
                        if hasattr(pf, 'alignment') and pf.alignment:
                            style_info['paragraph']['alignment'] = str(pf.alignment)
                        if hasattr(pf, 'line_spacing') and pf.line_spacing:
                            style_info['paragraph']['line_spacing'] = pf.line_spacing
                        if hasattr(pf, 'space_before') and pf.space_before:
                            style_info['paragraph']['space_before'] = pf.space_before
                        if hasattr(pf, 'space_after') and pf.space_after:
                            style_info['paragraph']['space_after'] = pf.space_after
                        if hasattr(pf, 'first_line_indent') and pf.first_line_indent:
                            style_info['paragraph']['first_line_indent'] = pf.first_line_indent
                    
                    styles[style.name] = style_info
            except Exception as e:
                logger.warning(f"提取样式 {style_name} 时出错: {str(e)}")
                continue
        
        # 确保基本样式存在
        if 'Normal' not in styles:
            styles['Normal'] = {
                'name': 'Normal',
                'font': {'name': 'Times New Roman', 'size': 12},
                'paragraph': {}
            }
        
        for i in range(1, 7):
            heading_name = f'Heading {i}'
            if heading_name not in styles:
                styles[heading_name] = {
                    'name': heading_name,
                    'font': {'name': 'Arial', 'size': 14 - i, 'bold': True},
                    'paragraph': {'space_before': 12, 'space_after': 6}
                }
        
        return styles
    
    def _extract_structure(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """
        提取文档结构
        
        Args:
            doc: docx文档对象
            
        Returns:
            结构元素列表
        """
        structure = []
        
        # 处理段落
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue  # 跳过空段落
                
            para_info = {
                'type': 'paragraph',
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else 'Normal'
            }
            
            # 检查是否为标题
            if para.style and para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.split(' ')[1])
                    para_info['type'] = 'heading'
                    para_info['level'] = level
                except (IndexError, ValueError):
                    pass
            
            structure.append(para_info)
        
        # 处理表格
        for i, table in enumerate(doc.tables):
            table_info = {
                'type': 'table',
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': []
            }
            
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    cell_info = {
                        'row': r,
                        'col': c,
                        'text': cell.text,
                        'style': cell.paragraphs[0].style.name if cell.paragraphs and cell.paragraphs[0].style else 'Normal'
                    }
                    table_info['cells'].append(cell_info)
            
            structure.append(table_info)
        
        return structure
    
    def _extract_page_settings(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        提取页面设置
        
        Args:
            doc: docx文档对象
            
        Returns:
            页面设置字典
        """
        settings = {}
        
        # 提取页面大小和边距
        for section in doc.sections:
            settings['page_width'] = section.page_width
            settings['page_height'] = section.page_height
            settings['left_margin'] = section.left_margin
            settings['right_margin'] = section.right_margin
            settings['top_margin'] = section.top_margin
            settings['bottom_margin'] = section.bottom_margin
            settings['header_distance'] = section.header_distance
            settings['footer_distance'] = section.footer_distance
            break  # 只处理第一个section
        
        return settings


class TexTemplateParser(TemplateParser):
    """
    .tex格式模板解析器
    """
    
    def parse(self, template_file: str) -> Dict[str, Any]:
        """
        解析.tex模板文件
        
        Args:
            template_file: .tex模板文件路径
            
        Returns:
            结构化的模板表示
        """
        logger.info(f"解析.tex模板: {template_file}")
        
        try:
            with open(template_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 提取文档类和包
            document_class = self._extract_document_class(content)
            packages = self._extract_packages(content)
            
            # 提取样式信息
            styles = self._extract_styles(content)
            
            # 提取文档结构
            structure = self._extract_structure(content)
            
            template_info = {
                'type': 'tex',
                'path': template_file,
                'document_class': document_class,
                'packages': packages,
                'styles': styles,
                'structure': structure
            }
            
            logger.info(f"解析完成，提取了 {len(packages)} 个包和 {len(structure)} 个结构元素")
            return template_info
            
        except Exception as e:
            logger.error(f"解析.tex模板时出错: {str(e)}", exc_info=True)
            raise
    
    def _extract_document_class(self, content: str) -> Dict[str, Any]:
        """
        提取文档类定义
        
        Args:
            content: 模板内容
            
        Returns:
            文档类信息
        """
        document_class = {'name': '', 'options': []}
        
        # 匹配文档类定义
        pattern = r'\\documentclass(?:\[(.*?)\])?\{(.*?)\}'
        match = re.search(pattern, content)
        if match:
            options_str = match.group(1)
            class_name = match.group(2)
            
            document_class['name'] = class_name
            if options_str:
                document_class['options'] = [opt.strip() for opt in options_str.split(',')]
        
        return document_class
    
    def _extract_packages(self, content: str) -> List[Dict[str, Any]]:
        """
        提取包引用
        
        Args:
            content: 模板内容
            
        Returns:
            包信息列表
        """
        packages = []
        
        # 匹配包引用
        pattern = r'\\usepackage(?:\[(.*?)\])?\{(.*?)\}'
        matches = re.finditer(pattern, content)
        
        for match in matches:
            options_str = match.group(1)
            package_name = match.group(2)
            
            package_info = {
                'name': package_name,
                'options': []
            }
            
            if options_str:
                package_info['options'] = [opt.strip() for opt in options_str.split(',')]
            
            packages.append(package_info)
        
        return packages
    
    def _extract_styles(self, content: str) -> Dict[str, Dict[str, Any]]:
        """
        提取样式定义
        
        Args:
            content: 模板内容
            
        Returns:
            样式定义字典
        """
        styles = {}
        
        # 提取标题样式
        heading_commands = ['section', 'subsection', 'subsubsection', 'paragraph', 'subparagraph']
        for cmd in heading_commands:
            # 检查是否有自定义标题格式
            pattern = r'\\renewcommand{\\' + cmd + r'}{(.*?)}'
            match = re.search(pattern, content)
            if match:
                styles[cmd] = {
                    'type': 'heading',
                    'definition': match.group(1)
                }
            else:
                styles[cmd] = {
                    'type': 'heading',
                    'definition': f'\\{cmd}'
                }
        
        # 提取环境样式
        environments = ['itemize', 'enumerate', 'description', 'quote', 'verbatim', 'tabular']
        for env in environments:
            # 检查是否有自定义环境
            pattern = r'\\renewenvironment{' + env + r'}{(.*?)}{(.*?)}'
            match = re.search(pattern, content)
            if match:
                styles[env] = {
                    'type': 'environment',
                    'begin_def': match.group(1),
                    'end_def': match.group(2)
                }
            else:
                styles[env] = {
                    'type': 'environment',
                    'begin_def': f'\\begin{{{env}}}',
                    'end_def': f'\\end{{{env}}}'
                }
        
        return styles
    
    def _extract_structure(self, content: str) -> List[Dict[str, Any]]:
        """
        提取文档结构
        
        Args:
            content: 模板内容
            
        Returns:
            结构元素列表
        """
        structure = []
        
        # 提取文档主体部分
        body_pattern = r'\\begin{document}(.*?)\\end{document}'
        body_match = re.search(body_pattern, content, re.DOTALL)
        if not body_match:
            return structure
        
        body_content = body_match.group(1)
        
        # 提取章节标题
        section_pattern = r'\\(section|subsection|subsubsection|paragraph|subparagraph)\{(.*?)\}'
        section_matches = re.finditer(section_pattern, body_content)
        
        for match in section_matches:
            section_type = match.group(1)
            section_title = match.group(2)
            
            level_map = {
                'section': 1,
                'subsection': 2,
                'subsubsection': 3,
                'paragraph': 4,
                'subparagraph': 5
            }
            
            structure.append({
                'type': 'heading',
                'level': level_map.get(section_type, 1),
                'text': section_title,
                'command': f'\\{section_type}'
            })
        
        # 提取环境
        env_pattern = r'\\begin\{(.*?)\}(.*?)\\end\{\1\}'
        env_matches = re.finditer(env_pattern, body_content, re.DOTALL)
        
        for match in env_matches:
            env_type = match.group(1)
            env_content = match.group(2)
            
            structure.append({
                'type': 'environment',
                'env_type': env_type,
                'content': env_content.strip()
            })
        
        return structure


class TemplateParserFactory:
    """
    模板解析器工厂，根据文件格式创建相应的解析器
    """
    
    @staticmethod
    def create_parser(format_type: str) -> TemplateParser:
        """
        创建模板解析器
        
        Args:
            format_type: 模板格式类型
            
        Returns:
            对应格式的模板解析器
        """
        if format_type.lower() == 'docx':
            return DocxTemplateParser()
        elif format_type.lower() == 'tex':
            return TexTemplateParser()
        else:
            raise ValueError(f"不支持的模板格式: {format_type}")
