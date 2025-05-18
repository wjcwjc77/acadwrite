"""
测试脚本，用于验证markdown到模板的映射功能
"""

import os
import sys
import argparse
import logging
import traceback

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.main import TemplateMapper

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def create_sample_markdown():
    """创建示例Markdown文件用于测试"""
    sample_md = """# 人工智能在医疗领域的应用综述

## 摘要

本文综述了人工智能技术在医疗领域的最新应用进展。随着深度学习、自然语言处理和计算机视觉等技术的发展，AI在医学影像分析、疾病诊断、药物研发和医疗管理等方面展现出巨大潜力。本文分析了当前研究现状、关键技术、应用案例以及面临的挑战与未来发展方向。

## 1. 引言

人工智能(AI)技术在过去十年中取得了突破性进展，其在医疗健康领域的应用也日益广泛。医疗AI系统能够处理和分析大量医疗数据，辅助医生进行诊断决策，提高医疗效率和准确性。本文旨在全面回顾AI在医疗领域的应用现状，并探讨其未来发展趋势。

## 2. 研究方法

本综述采用系统文献回顾方法，检索了2015-2025年间发表的相关研究文献。我们使用Web of Science、PubMed和IEEE Xplore等数据库，以"医疗人工智能"、"深度学习医疗应用"等关键词进行检索，最终纳入分析的文献共计235篇。

## 3. 关键技术

### 3.1 深度学习

深度学习是当前医疗AI应用的核心技术，特别是卷积神经网络(CNN)在医学影像分析中表现出色。例如，ResNet和U-Net等网络架构被广泛应用于肿瘤检测和器官分割任务。

### 3.2 自然语言处理

NLP技术能够从电子病历、医学文献和临床笔记中提取有价值的信息。BERT和GPT等预训练模型在医学文本理解和生成方面取得了显著进展。

### 3.3 强化学习

强化学习在个性化治疗方案制定和药物剂量优化等方面展现出潜力，通过不断学习和调整决策来优化治疗效果。

## 4. 应用领域

### 4.1 医学影像分析

AI在放射学、病理学和皮肤科等影像分析中的应用最为成熟。例如，DeepMind开发的眼底图像分析系统可以检测50多种眼部疾病，准确率达到专科医生水平。

### 4.2 疾病诊断与预测

基于机器学习的诊断系统能够整合患者的多维度数据，预测疾病风险和发展趋势。例如，Mayo Clinic开发的AI系统可以提前预测心脏病发作风险。

### 4.3 药物研发

AI加速了新药发现和开发过程。例如，Insilico Medicine利用生成对抗网络设计的新分子，从发现到临床前测试仅用了不到18个月。

### 4.4 智能医疗管理

AI在医院管理、患者流程优化和医疗资源分配等方面也发挥着重要作用，提高医疗系统的整体效率。

## 5. 挑战与局限性

尽管AI在医疗领域取得了显著进展，但仍面临数据隐私、算法透明度、临床验证和伦理问题等多重挑战。特别是，如何确保AI系统的公平性和可解释性是当前研究的重点。

## 6. 未来发展趋势

未来医疗AI将朝着多模态融合、联邦学习、自监督学习和人机协作等方向发展。随着技术进步和监管框架的完善，AI将在更广泛的医疗场景中发挥作用。

## 7. 结论

人工智能正在深刻变革医疗健康领域，为提高诊断准确性、治疗效果和医疗可及性带来新的可能。未来需要医学专家、AI研究者和政策制定者的紧密合作，共同推动医疗AI的负责任发展和应用。

## 参考文献

1. Smith J, et al. (2023). Deep learning applications in medical imaging: A comprehensive review. Nature Medicine, 29(3), 456-470.
2. Wang L, et al. (2022). Natural language processing for electronic health records: Progress and challenges. JAMA, 327(5), 378-388.
3. Chen T, et al. (2024). Reinforcement learning for personalized treatment planning in oncology. Science Translational Medicine, 16(4), eabc1234.
"""
    
    sample_md_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'sample_article.md')
    with open(sample_md_path, 'w', encoding='utf-8') as f:
        f.write(sample_md)
    
    return sample_md_path


def create_sample_docx_template():
    """创建示例.docx模板文件用于测试"""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'sample_template.docx')
    
    doc = docx.Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加标题样式
    styles = doc.styles
    
    # 修改标题1样式
    style_heading1 = styles['Heading 1']
    style_heading1.font.name = 'Arial'
    style_heading1.font.size = Pt(16)
    style_heading1.font.bold = True
    style_heading1.paragraph_format.space_before = Pt(12)
    style_heading1.paragraph_format.space_after = Pt(6)
    
    # 修改标题2样式
    style_heading2 = styles['Heading 2']
    style_heading2.font.name = 'Arial'
    style_heading2.font.size = Pt(14)
    style_heading2.font.bold = True
    style_heading2.paragraph_format.space_before = Pt(10)
    style_heading2.paragraph_format.space_after = Pt(4)
    
    # 修改标题3样式
    style_heading3 = styles['Heading 3']
    style_heading3.font.name = 'Arial'
    style_heading3.font.size = Pt(12)
    style_heading3.font.bold = True
    style_heading3.paragraph_format.space_before = Pt(8)
    style_heading3.paragraph_format.space_after = Pt(4)
    
    # 修改正文样式
    style_normal = styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15
    
    # 添加示例内容
    doc.add_paragraph('示例文档标题', style='Heading 1')
    doc.add_paragraph('这是一个示例段落，用于展示文档样式。', style='Normal')
    doc.add_paragraph('第一章', style='Heading 2')
    doc.add_paragraph('这是第一章的内容，包含了一些示例文本。', style='Normal')
    doc.add_paragraph('第一节', style='Heading 3')
    doc.add_paragraph('这是第一节的详细内容，用于测试样式应用。', style='Normal')
    
    # 保存文档
    doc.save(template_path)
    
    return template_path


def create_sample_tex_template():
    """创建示例.tex模板文件用于测试"""
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'sample_template.tex')
    
    template_content = r"""\documentclass[12pt,a4paper]{article}

\usepackage{geometry}
\usepackage{titlesec}
\usepackage{setspace}
\usepackage{graphicx}
\usepackage{hyperref}

% 设置页面边距
\geometry{top=2.5cm, bottom=2.5cm, left=2.5cm, right=2.5cm}

% 设置标题格式
\titleformat{\section}{\Large\bfseries}{\thesection}{1em}{}
\titleformat{\subsection}{\large\bfseries}{\thesubsection}{1em}{}
\titleformat{\subsubsection}{\normalsize\bfseries}{\thesubsubsection}{1em}{}

% 设置行距
\onehalfspacing

\begin{document}

\title{示例文档标题}
\author{作者名}
\date{\today}
\maketitle

\section{第一章}
这是一个示例段落，用于展示文档样式。

\subsection{第一节}
这是第一节的内容，包含了一些示例文本。

\subsubsection{小节}
这是小节的详细内容，用于测试样式应用。

\end{document}
"""
    
    with open(template_path, 'w', encoding='utf-8') as f:
        f.write(template_content)
    
    return template_path


def test_docx_mapping():
    """测试Markdown到.docx模板的映射"""
    logger.info("开始测试Markdown到.docx模板的映射")
    
    # 创建示例文件
    markdown_file = create_sample_markdown()
    template_file = create_sample_docx_template()
    
    # 设置输出文件路径
    output_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output_docx.docx')
    
    try:
        # 创建映射器并处理
        mapper = TemplateMapper()
        result = mapper.process(markdown_file, template_file, output_file)
        
        logger.info(f"映射完成，输出文件: {result}")
        logger.info("请手动检查输出文件的格式和样式是否符合预期")
        
        return True
    except Exception as e:
        logger.error(f"测试失败: {str(e)}")
        logger.error(traceback.format_exc())  # 添加详细的异常堆栈信息
        return False


def test_tex_mapping():
    """测试Markdown到.tex模板的映射"""
    logger.info("开始测试Markdown到.tex模板的映射")
    
    # 创建示例文件
    markdown_file = create_sample_markdown()
    template_file = create_sample_tex_template()
    
    # 设置输出文件路径
    output_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output_tex.tex')
    
    try:
        # 创建映射器并处理
        mapper = TemplateMapper()
        result = mapper.process(markdown_file, template_file, output_file)
        
        logger.info(f"映射完成，输出文件: {result}")
        logger.info("请手动检查输出文件的格式和样式是否符合预期")
        
        return True
    except Exception as e:
        logger.error(f"测试失败: {str(e)}")
        logger.error(traceback.format_exc())  # 添加详细的异常堆栈信息
        return False


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='测试Markdown到模板的映射功能')
    parser.add_argument('--format', choices=['docx', 'tex', 'all'], default='all',
                        help='要测试的模板格式')
    
    args = parser.parse_args()
    
    if args.format == 'docx' or args.format == 'all':
        docx_result = test_docx_mapping()
        print(f"Markdown到.docx模板映射测试: {'成功' if docx_result else '失败'}")
    
    if args.format == 'tex' or args.format == 'all':
        tex_result = test_tex_mapping()
        print(f"Markdown到.tex模板映射测试: {'成功' if tex_result else '失败'}")


if __name__ == "__main__":
    main()
