import docx
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm
from docx.enum.text import *
import os
import sys
from docx import Document
from PyQt5.QtWidgets import QApplication, QFileDialog
# 删除段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
# p._p = p._element = None
    paragraph._p = paragraph._element = None
 
#判断是否为落款格式
def LuoKuan(str):
    for i in str:
        if i in punc:
            return False
    if ((str[0] in num) and (str[-1] == "日") and (len(str) <= 12)) or ((str[0] in cn_num) and (str[-1] == "日") and (len(str) <= 12)):
        return True
    else:
        return False
def setMargin(docx):
    section = docx.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
 
#判断是否为一级标题格式（如：一、xxx）
def GradeOneTitle(str):
    if ((str[0] in cn_num) and (str[1] == "、")) or ((str[0] in cn_num) and (str[1] in cn_num) and (str[2] == "、")):
        return True
    else:
        return False
 
#判断是否为二级标题格式（如：（一）xxx）
def GradeTwoTitle(str):
    if ((str[0] == "（") and (str[1] in cn_num) and (str[2] == "）")) or ((str[0] == "（") and (str[1] in cn_num) and (str[2] in cn_num) and (str[3] == "）")):
        return True
    else:
        return False
#判断是否为三级标题格式（如：1.xxx）
def GradeThreeTitle(str):
    if ((str[0] in num) and (str[1] in punc)) or ((str[0] in num) and (str[1] in num) and (str[2] in punc)):
        return True
    else:
        return False
#判断是否为四级标题格式（如：（1）xxx）
def GradeFourTitle(str):
    if ((str[0] == "（") and (str[1] in num) and (str[2] == "）")) or ((str[0] == "（") and (str[1] in num) and (str[2] in num) and (str[3] == "）")):
        return True
    else:
        return False
 
#判断是否为五级标题格式（如：一是XXX）
def GradeFiveTitle(str):
    if ((str[0] in cn_num) and (str[1] in must)) or ((str[0] in cn_num) and (str[1] in cn_num) and (str[1] in must)):
        return True
    else:
        return False
 
def OneKeyWord():
    global cn_num,num,punc,must
    cn_num = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    num = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "0"]
    punc = ["。", "，", "！", "？", "：", "；", "、", ".", "（", "）","．"]
    must = ["要", "是", "能"]
    filecnt = 0
    print('欢迎使用Word一键排版工具！创作者QQ：124500535')
    confirm= input("是否打开Ｗord文档？输入“Y”表示“打开”，输入“N”表示“取消”！")
    if confirm == 'Y' or confirm == 'y':
        a = QApplication([''])
        files, stylel = QFileDialog.getOpenFileNames(caption="多文件选择", directory="/", filter="Word 文档(*.docx)")
        print(files)  # 打印所选文件全部路径（包括文件名和后缀名）和文件类型
        for file in files:
            docx = Document(file)
            paragraphcnt = 0
            filecnt= filecnt+1
            print('这是第%s个文件：%s' %(filecnt,file))
            for paragraph in docx.paragraphs:
                paragraphcnt = paragraphcnt +1
                paragraph.text=paragraph.text.replace(",","，")
                paragraph.text=paragraph.text.replace(";","；")
                paragraph.text=paragraph.text.replace(":","：")
                paragraph.text=paragraph.text.replace("!","！")
                paragraph.text=paragraph.text.replace("?","？")
                paragraph.text=paragraph.text.replace("(","（")
                paragraph.text=paragraph.text.replace(")","）")
                paragraph.text=paragraph.text.replace(" ","")
                paragraph.text=paragraph.text.replace("\t", "")
                paragraph.text = paragraph.text.replace("\n", "")
                if paragraph.text == '':
                    delete_paragraph(paragraph)
                    paragraphcnt = paragraphcnt-1
                    continue
                paragraph.paragraph_format.left_indent = 0  #预先对缩进赋值, 防止对象为空报错
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '0')  #并去除缩进
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLine"), '0')
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:leftChars"), '0')
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:left"), '0')
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:rightChars"), '0')
                paragraph.paragraph_format.element.pPr.ind.set(qn("w:right"), '0')
                print('这是第%s段' %paragraphcnt)
                print(paragraph.text)
                if paragraphcnt == 1 and len(paragraph.text)<40:
                    #处理头部空行
                    #标题（方正小标宋_GBK、2号、加粗、居中、下端按2号字空一行）
                    paragraph.paragraph_format.line_spacing=Pt(28)  #行距固定值28磅
                    paragraph.paragraph_format.space_after = Pt(0)  #段后间距=0
                    for run in paragraph.runs:
                        run.font.size = Pt(22)  # 字体大小2号
                        run.bold = False  # 加粗
                        run.font.name = '方正小标宋_GBK'  # 控制是西文时的字体
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')  # 控制是中文时的字体
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中
                    continue
                elif paragraphcnt == 2 and len(paragraph.text) < 30:
                    # 作者单位、姓名
                    paragraph.paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                    paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
                    for run in paragraph.runs:
                        run.font.size = Pt(16)  # 字体大小2号
                        run.bold = False  # 加粗
                        run.font.name = '楷体'  # 控制是西文时的字体
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')  # 控制是中文时的字体
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
                    continue
                elif paragraphcnt == 3 and len(paragraph.text) < 30 and (paragraph.text[0] == "（") and (paragraph.text[1] in num):
                    # 日期，如（2023年6月15日）
                    paragraph.paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                    paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
                    for run in paragraph.runs:
                        run.font.size = Pt(16)  # 字体大小2号
                        run.bold = False  # 加粗
                        run.font.name = '楷体'  # 控制是西文时的字体
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')  # 控制是中文时的字体
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
                    continue
                    # #处理正文
                else:
                    paragraph.paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                    paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
                    paragraph.paragraph_format.first_line_indent = Pt(32)
                    for run in paragraph.runs:
                        run.font.size = Pt(16)  # 字体大小3号
                        run.bold = False  # 字体不加粗
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        if GradeOneTitle(run.text): #判断是否为一级标题格式（如：一、xxx）
                            run.font.name = '黑体'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        elif GradeTwoTitle(run.text): #判断是否为二级标题格式（如：（一）xxx）
                            if "。" not in run.text:
                                run.font.name = '楷体'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                            else:
                                run.text = run.text.split('。',1)
                                run.font.name = '楷体'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                        elif GradeThreeTitle(run.text): #判断是否为三级标题格式（如：1.xxx）
                            if "。" not in run.text:
                                if (run.text[0] in num) and (run.text[1] in punc):
                                    run.text = run.text.replace(run.text[1], "．",1)
                                if (run.text[0] in num) and (run.text[1] in num) and (run.text[2] in punc):
                                    run.text = run.text.replace(run.text[2], "．", 1)
                                run.font.name = '仿宋_GB2312'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                run.bold = True  # 字体加粗
                            else:
                                if (run.text[0] in num) and (run.text[1] in punc):
                                    run.text = run.text.replace(run.text[1], "．", 1)
                                if (run.text[0] in num) and (run.text[1] in num) and (run.text[2] in punc):
                                    run.text = run.text.replace(run.text[2], "．", 1)
                                sentence_to_bold = run.text.split('。')[0]+"。"
                                sentence_not_to_bold = run.text.split('。',1)[1]
                                paragraph.insert_paragraph_before(sentence_to_bold)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.first_line_indent = Pt(32)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.space_after = Pt(0)  # 段后间距=0
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                docx.paragraphs[paragraphcnt - 1].runs[0].bold = True  # 字体加粗
                                docx.paragraphs[paragraphcnt - 1].add_run(sentence_not_to_bold).bold = False
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[1].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                delete_paragraph(paragraph)
                        elif GradeFourTitle(run.text): #判断是否为四级标题格式（如：（1）xxx）
                            if "。" not in run.text:
                                run.font.name = '仿宋_GB2312'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                run.bold = True  # 字体加粗
                            else:
                                sentence_to_bold = run.text.split('。')[0]+"。"
                                sentence_not_to_bold = run.text.split('。',1)[1]
                                paragraph.insert_paragraph_before(sentence_to_bold)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.first_line_indent = Pt(32)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.space_after = Pt(0)  # 段后间距=0
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                docx.paragraphs[paragraphcnt - 1].runs[0].bold = True  # 字体加粗
                                docx.paragraphs[paragraphcnt - 1].add_run(sentence_not_to_bold).bold = False
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[1].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                delete_paragraph(paragraph)
                        elif GradeFiveTitle(run.text): #判断是否为五级标题格式（如：一是xxx）
                            if "。" not in run.text:                                
                                run.font.name = '仿宋_GB2312'
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                run.bold = True  # 字体加粗
                            else:                                
                                sentence_to_bold = run.text.split('。')[0]+"。"
                                sentence_not_to_bold = run.text.split('。',1)[1]
                                paragraph.insert_paragraph_before(sentence_to_bold)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.first_line_indent = Pt(32)
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.line_spacing = Pt(28)  # 行距固定值28磅
                                docx.paragraphs[paragraphcnt - 1].paragraph_format.space_after = Pt(0)  # 段后间距=0
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[0].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                docx.paragraphs[paragraphcnt - 1].runs[0].bold = True  # 字体加粗
                                docx.paragraphs[paragraphcnt - 1].add_run(sentence_not_to_bold).bold = False
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.name = '仿宋_GB2312'
                                docx.paragraphs[paragraphcnt - 1].runs[1].font.size = Pt(16)  # 字体大小3号
                                docx.paragraphs[paragraphcnt - 1].runs[1].element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                delete_paragraph(paragraph)
                        elif LuoKuan(run.text):  # 判断是否为落款格式
                            run.font.name = '仿宋_GB2312'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                            run.text = "\r" * 2 + run.text  # 前置空格，顶到最右，需手动调整空格
                            paragraph.paragraph_format.left_indent = Pt(288)    #18B*16Pt=288Pt
                        else: #普通正文格式
                            run.font.name = '仿宋_GB2312'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            setMargin(docx)
            docx.save(file)
if __name__ == '__main__':
    OneKeyWord()
    os.system("pause")